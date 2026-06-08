#!/usr/bin/env python3
"""Genera un CV profesional y honesto (.docx) + carta de presentacion
para postular a Tecnico Electronico junior (Springfield Electronica).

Solo quedan por completar los datos estrictamente personales:
NOMBRE, TELEFONO, EMAIL y BARRIO. Todo lo demas ya esta redactado.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)
NEGRO = RGBColor(0x00, 0x00, 0x00)
ROJO = RGBColor(0xB0, 0x00, 0x00)


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL
    run.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F4E79')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_bullet(doc, text, fill=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    if fill:
        run.font.color.rgb = GRIS
        run.italic = True
    return p


def add_field(doc, label, value, fill=False):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Calibri'
    if fill:
        r2.font.color.rgb = GRIS
        r2.italic = True
    return p


def base_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    return doc


# ============================================================
#                           CV
# ============================================================
doc = base_doc()

# ---- ENCABEZADO ----
nombre = doc.add_paragraph()
nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nombre.add_run("[NOMBRE Y APELLIDO]")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = AZUL

subt = doc.add_paragraph()
subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subt.add_run("Tecnico Electronico  |  Instalacion y mantenimiento de sistemas electronicos e informatica")
r.font.size = Pt(11)
r.font.color.rgb = GRIS

contacto = doc.add_paragraph()
contacto.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = contacto.add_run("[Telefono]   |   [Email]   |   [Barrio], CABA   |   24 anos")
r.font.size = Pt(10)
r.font.color.rgb = GRIS

# ---- PERFIL ----
add_heading(doc, "Perfil profesional")
p = doc.add_paragraph()
r = p.add_run(
    "Tecnico Electronico egresado de la E.T. N.\u00ba 1 \u201cOtto Krause\u201d, con base solida en "
    "electronica, redes e informatica. Busco desarrollarme como Tecnico de Campo en instalacion "
    "y mantenimiento de sistemas de videovigilancia, control de acceso e informatica general. "
    "Me destaco por mi responsabilidad, predisposicion al trabajo en equipo y una fuerte vocacion "
    "de aprendizaje continuo: estoy abierto a capacitarme y a sumar tareas en todas las areas. "
    "Disponibilidad horaria completa y movilidad para trabajar en clientes dentro de CABA."
)
r.font.size = Pt(10.5)
r.font.color.rgb = NEGRO

# ---- FORTALEZAS CLAVE ----
add_heading(doc, "Por que encajo en el puesto")
add_bullet(doc, "Tecnico graduado: cumplo el requisito de titulo (plan completo, todas las materias aprobadas).")
add_bullet(doc, "Perfil junior con muchas ganas de aprender; valoro la capacitacion que ofrece la empresa.")
add_bullet(doc, "Formacion tecnica orientada a las tres areas de la empresa: electronica, seguridad e informatica.")
add_bullet(doc, "Disponibilidad para tareas de campo y para desempenarme en todas las areas.")

# ---- FORMACION ----
add_heading(doc, "Formacion academica")
add_field(doc, "Titulo", "Tecnico Electronico \u2014 E.T. N.\u00ba 1 \u201cOtto Krause\u201d, CABA.")
add_field(doc, "Estado", "Egresado \u2014 plan de estudios completo, todas las materias aprobadas. Titulo en tramite "
                          "(disponible certificado analitico / constancia de titulo en tramite).")
add_field(doc, "Ano de egreso", "[completar ano]", fill=True)

# ---- COMPETENCIAS TECNICAS (por area de la empresa) ----
add_heading(doc, "Competencias tecnicas")
p = doc.add_paragraph(); r = p.add_run("Videovigilancia y seguridad electronica"); r.bold = True; r.font.size = Pt(10.5)
add_bullet(doc, "Nociones de camaras CCTV / IP, cableado y conexionado de sistemas de seguridad.")
add_bullet(doc, "Interes y predisposicion para capacitarme en control de acceso y alarmas.")
p = doc.add_paragraph(); r = p.add_run("Electronica"); r.bold = True; r.font.size = Pt(10.5)
add_bullet(doc, "Electronica analogica y digital; uso de instrumental (multimetro, osciloscopio, fuente).")
add_bullet(doc, "Soldadura, armado, medicion y diagnostico de circuitos y componentes.")
p = doc.add_paragraph(); r = p.add_run("Informatica y redes"); r.bold = True; r.font.size = Pt(10.5)
add_bullet(doc, "Armado, mantenimiento y reparacion de PC; instalacion de software y perifericos.")
add_bullet(doc, "Redes basicas (TCP/IP, cableado de red, configuracion de routers/switches).")

# ---- EXPERIENCIA / PROYECTOS ----
add_heading(doc, "Experiencia y proyectos")
p = doc.add_paragraph()
r = p.add_run("Completa cada punto con datos reales y borra los que no apliquen. "
              "(No agregues trabajos que no hiciste.)")
r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRIS
add_bullet(doc, "Proyecto final / trabajos practicos del Otto Krause: [nombre del proyecto y que hacia].", fill=True)
add_bullet(doc, "Armado y reparacion de PC para conocidos y vecinos: diagnostico, cambio de componentes, "
                "instalacion de sistema operativo. [ajustar]", fill=True)
add_bullet(doc, "Instalaciones / reparaciones por cuenta propia (redes hogarenas, camaras, etc.): [detallar].", fill=True)

# ---- HABILIDADES PERSONALES ----
add_heading(doc, "Habilidades personales")
add_bullet(doc, "Gran predisposicion y entusiasmo por aprender cosas nuevas.")
add_bullet(doc, "Responsabilidad, puntualidad y compromiso.")
add_bullet(doc, "Trabajo en equipo y buen trato con clientes.")
add_bullet(doc, "Resolucion de problemas y atencion al detalle.")

# ---- CURSOS ----
add_heading(doc, "Cursos (recomendados para sumar valor rapido)")
p = doc.add_paragraph()
r = p.add_run("Cursos cortos y gratuitos. Hacelos y agregalos aca con la fecha real (suman mucho para este puesto):")
r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRIS
add_bullet(doc, "CCTV / Videovigilancia IP \u2014 academias online de fabricantes (Hikvision, Dahua).", fill=True)
add_bullet(doc, "Redes y cableado estructurado \u2014 Cisco \u201cNetworking Basics\u201d (gratis).", fill=True)
add_bullet(doc, "Introduccion a la ciberseguridad \u2014 Cisco / Fundacion Telefonica.", fill=True)

# ---- OTROS DATOS ----
add_heading(doc, "Otros datos")
add_field(doc, "Disponibilidad", "Horaria completa; el horario del puesto (8:30 a 14:30 h) no representa inconveniente.")
add_field(doc, "Movilidad", "Disponibilidad para realizar trabajos en clientes dentro de CABA. [Indicar si tenes licencia de conducir]", fill=True)
add_field(doc, "Idiomas", "Espanol (nativo). Ingles: [nivel basico/intermedio \u2014 completar]", fill=True)
add_field(doc, "Informatica", "Windows, paquete Office, diagnostico y mantenimiento de equipos.")

# ---- NOTA ----
doc.add_paragraph()
nota = doc.add_paragraph()
r = nota.add_run(
    "--- NOTA (BORRAR ANTES DE ENVIAR): Solo falta completar lo que esta en gris/itálica y "
    "entre [corchetes] (nombre, contacto, barrio, ano de egreso, etc.). Guardalo como PDF para enviarlo. ---"
)
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = ROJO

doc.save("CV_Tecnico_Electronico.docx")


# ============================================================
#                  CARTA DE PRESENTACION
# ============================================================
carta = base_doc()

h = carta.add_paragraph()
r = h.add_run("[NOMBRE Y APELLIDO]")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = AZUL
sc = carta.add_paragraph()
r = sc.add_run("[Telefono]  |  [Email]  |  [Barrio], CABA")
r.font.size = Pt(10); r.font.color.rgb = GRIS

carta.add_paragraph()
dest = carta.add_paragraph()
r = dest.add_run("A: Springfield Electronica  \u2014  info@springfield.com.ar")
r.bold = True; r.font.size = Pt(10.5)
asunto = carta.add_paragraph()
r = asunto.add_run("Asunto: Postulacion \u2014 Tecnico Electronico / Computacion (junior)")
r.bold = True; r.font.size = Pt(10.5)
carta.add_paragraph()

parrafos = [
    "Estimados:",
    "Me dirijo a ustedes para postularme al puesto de Tecnico Electronico / en Computacion de "
    "nivel junior. Soy egresado de la E.T. N.\u00ba 1 \u201cOtto Krause\u201d como Tecnico Electronico, "
    "con el plan de estudios completo y todas las materias aprobadas; mi titulo se encuentra en "
    "tramite y cuento con el certificado analitico para acreditarlo.",
    "Si bien busco mi primera experiencia laboral formal, tengo una base tecnica solida en "
    "electronica, redes e informatica y muchas ganas de aprender. Me interesa especialmente que el "
    "puesto incluya capacitacion y trabajo de campo en videovigilancia, control de acceso e "
    "informatica, areas en las que quiero desarrollarme. Estoy dispuesto a desempenarme en todas "
    "las areas que maneja la empresa.",
    "Cuento con disponibilidad horaria para el turno de 8:30 a 14:30 h y disponibilidad para "
    "realizar trabajos en clientes dentro de CABA. Adjunto mi CV y quedo a disposicion para una "
    "entrevista.",
    "Desde ya, muchas gracias por su tiempo.",
    "Saludos cordiales,",
    "[NOMBRE Y APELLIDO]",
]
for i, t in enumerate(parrafos):
    p = carta.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(10.5)
    if t.startswith("[NOMBRE"):
        r.bold = True

carta.save("Carta_Presentacion.docx")

print("Generados: CV_Tecnico_Electronico.docx y Carta_Presentacion.docx")
